from flask import Flask, render_template, request, send_file
import cv2
import numpy as np
from PIL import Image,ImageFilter,ImageEnhance
import pytesseract
pytesseract.pytesseract.tesseract_cmd = '/usr/bin/tesseract'
from docx import Document
from datetime import datetime
import os

def extract_text(image_path, language):
    # Load the image from file
    image = Image.open(image_path)
    image_e = image.convert('L')  # 转换为灰度图像
    image_e = image_e.filter(ImageFilter.MedianFilter())  # 中值滤波去噪
    image_e = ImageEnhance.Contrast(image_e).enhance(10)  # 增强对比度
    # image_e = image_e.point(lambda x: 0 if x < 140 else 255)
    # Use tesseract to do OCR on the image
    custom_config = r'--oem 1'
    try:
        text = pytesseract.image_to_string(image_e, lang=language, config=custom_config)
    except pytesseract.TesseractError as e:
        print("An error occurred during OCR processing:", e)
        return None

    return text


def extract_subplots(image_path):
    image = Image.open(image_path)
    gray = cv2.cvtColor(np.array(image), cv2.COLOR_BGR2GRAY)
    _, binary = cv2.threshold(gray, 128, 255, cv2.THRESH_BINARY_INV)
    contours, _ = cv2.findContours(binary, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

    def is_image_contour(contour, image_shape):
        x, y, w, h = cv2.boundingRect(contour)
        aspect_ratio = w / float(h)
        area = w * h
        image_area = image_shape[0] * image_shape[1]
        return 0.25 < aspect_ratio < 4 and area > 0.01 * image_area

    image_contours = [cnt for cnt in contours if is_image_contour(cnt, gray.shape)]
    image_contours.sort(key=lambda cnt: cv2.boundingRect(cnt)[1])
    subplot_image_paths = []

    for i, contour in enumerate(image_contours):
        x, y, w, h = cv2.boundingRect(contour)
        subplot_image = image.crop((x, y, x + w, y + h))
        subplot_image_path = f'subplot_{i}.png'
        subplot_image.save(subplot_image_path)
        subplot_image_paths.append(subplot_image_path)

    return subplot_image_paths


# def create_markdown(text, subplot_image_paths, save_path):
#     markdown_content = ""
#     if text:
#         markdown_content += text + "\n\n"
#     for image_path in subplot_image_paths:
#         markdown_content += f"![Subplot Image]({image_path})\n\n"
#     with open(save_path, 'w', encoding='utf-8') as md_file:
#         md_file.write(markdown_content)


def create_docx(text, subplot_image_paths, save_path):
    doc = Document()
    if text:
        doc.add_paragraph(text)
    for image_path in subplot_image_paths:
        # sub_image=Image.open(image_path)
        doc.add_picture(image_path)
    #    docx_io = io.BytesIO()
    doc.save(save_path)


def save_image_to_disk(image):
    # 定义保存图片的路径
    save_path = "uploaded_images"
    if not os.path.exists(save_path):
        os.makedirs(save_path)

    # 定义图片文件名（这里简单使用一个计数器或时间戳来命名）
    image_filename = os.path.join(save_path, str(datetime.now().second)+"image.jpg")

    # 保存图片
    image.save(image_filename)

    return image_filename

def convert_pic(image, file_type, text_language):
    image_path = save_image_to_disk(image)
    extracted_text = extract_text(image_path, language=text_language)
    extracted_text = extracted_text.replace(' ', '')
    extracted_text = extracted_text.replace('\n', '')
    subplot_image_paths = extract_subplots(image_path)
    save_path = str.split(image_path, '.')[0] + '.' + file_type
    # if file_type == 'md':
    #     create_markdown(extracted_text, subplot_image_paths, save_path)
    if file_type == 'docx':
        create_docx(extracted_text, subplot_image_paths, save_path)
    return save_path

app = Flask(__name__)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload():
    file = request.files['image']
    image = Image.open(file.stream)
    docx_filename=convert_pic(image,'docx','chi_sim')
    return send_file(docx_filename, as_attachment=True)

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000)
